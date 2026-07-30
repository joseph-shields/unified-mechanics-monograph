"""Build THE FIELD LEDGER, second edition, in the corpus house style.

House primitives measured from the corpus. Additions over the A04 builder:

  raw=True on para and _cell, so a file name keeps its underscores instead of
    losing them to the subscript parser.
  cantSplit on callouts, without which a one-cell callout breaks across a page
    and orphans its last block on a sheet of its own.
  grade(), which prints the Ledger's status boxes as a shaded strip, since the
    grade is part of the argument and not an aside.
  part() and chap(), the book's two levels of division.
"""
import os, json, math
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FOLD = r"C:\Users\joesh\Desktop\SCI\Field Ledger Second Edition"
OUT = os.path.join(FOLD, "Unified_Mechanics_The_Field_Ledger_Second_Edition.docx")
IMG = os.path.join(FOLD, "figures")

SERIF, MATH = "Constantia", "Cambria Math"
NAVY = RGBColor(0x10, 0x2A, 0x43); GOLD = RGBColor(0xB4, 0x8A, 0x2E)
SLATE = RGBColor(0x5B, 0x6B, 0x7A); INK = RGBColor(0x24, 0x30, 0x40)
TEAL = RGBColor(0x14, 0x70, 0x7D); RUST = RGBColor(0xA8, 0x44, 0x2A)
MAUVE = RGBColor(0x7B, 0x5A, 0x80)
HEADFILL, BOXFILL, RULECOL = "DDE5EC", "EEF2F5", "B48A2E"
TEXT_W = 6.94

# the grade palette: one fill per status, so a reader can find a grade by colour
GRADEFILL = {
    "STANDARD":    ("EDF1F4", SLATE),
    "PROVED":      ("E4EFF0", TEAL),
    "CONDITIONAL": ("F6F1E4", GOLD),
    "PROPOSED":    ("F6EFF3", RGBColor(0xB8, 0x25, 0x6B)),
    "FORECAST":    ("EAF0F6", RGBColor(0x1E, 0x6F, 0xB8)),
    "OPEN":        ("F2F4F6", NAVY),
    "WITHDRAWN":   ("F2EEF3", MAUVE),
    "READING":     ("F4F1E8", GOLD),
}

doc = Document()
s = doc.sections[0]
s.page_width, s.page_height = Inches(8.5), Inches(11)
s.left_margin = s.right_margin = Inches(0.78)
s.top_margin = s.bottom_margin = Inches(0.68)
s.header_distance = s.footer_distance = Inches(0.38)
n = doc.styles["Normal"]; n.font.name = SERIF; n.font.size = Pt(10.5)
n.font.color.rgb = INK; n.paragraph_format.space_after = Pt(0)
n._element.rPr.rFonts.set(qn("w:eastAsia"), SERIF)


def _scripts(t):
    out, buf, i = [], "", 0
    while i < len(t):
        if t[i] in "_^" and i + 1 < len(t):
            m = "sub" if t[i] == "_" else "sup"; i += 1
            if t[i] == "{":
                j = t.index("}", i); tok, i = t[i+1:j], j+1
            else:
                tok, i = t[i], i+1
            if buf: out.append((buf, None)); buf = ""
            out.append((tok, m))
        else:
            buf += t[i]; i += 1
    if buf: out.append((buf, None))
    return out


def _runs(p, text, size, color, font, bold=False, italic=False, caps=False,
          raw=False):
    for i, ch in enumerate(text.split("**")):
        if not ch: continue
        for sub, m in ([(ch, None)] if raw else _scripts(ch)):
            r = p.add_run(sub); r.font.name = font; r.font.size = Pt(size)
            r.font.color.rgb = color; r.bold = bold or (i % 2 == 1)
            r.italic = italic
            if m == "sub": r.font.subscript = True
            elif m == "sup": r.font.superscript = True
            if caps: r.font.all_caps = True
            r._element.rPr.rFonts.set(qn("w:eastAsia"), font)


def para(text="", size=10.5, color=INK, bold=False, italic=False, font=SERIF,
         align=WD_ALIGN_PARAGRAPH.JUSTIFY, before=0, after=6, line=1.20,
         indent=0, keep=False, caps=False, brk=False, raw=False, style=None):
    p = doc.add_paragraph()
    # The style is set FIRST so that every direct format below overrides it. The
    # only thing carried over from the built-in heading is its outline level,
    # which is what Word's Navigation Pane, a generated table of contents and
    # the exported PDF bookmarks all read. The look is unchanged.
    if style:
        p.style = doc.styles[style]
    pf = p.paragraph_format
    pf.alignment = align; pf.space_before = Pt(before); pf.space_after = Pt(after)
    pf.line_spacing = line; pf.keep_with_next = keep
    if indent: pf.left_indent = Inches(indent)
    if brk: pf.page_break_before = True
    _runs(p, text, size, color, font, bold, italic, caps, raw)
    return p


def _border(p, color=RULECOL, sz=6, space=4, edge="bottom"):
    pPr = p._p.get_or_add_pPr(); b = OxmlElement("w:pBdr")
    e = OxmlElement("w:" + edge)
    e.set(qn("w:val"), "single"); e.set(qn("w:sz"), str(sz))
    e.set(qn("w:space"), str(space)); e.set(qn("w:color"), color)
    b.append(e); pPr.append(b)


def part(num, title, blurb):
    """A part opener: its own page, set large and quiet.

    The break goes on the SPACER, not on a paragraph of its own before it. With
    two paragraphs here, a preceding chapter that happens to end exactly at a
    page boundary pushes the first one onto a sheet by itself and the book gains
    an entirely blank page.
    """
    para("", size=1, after=150, brk=True)
    FIRST_OF_PART[0] = True
    para(num, size=10, color=GOLD, bold=True, caps=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, after=10)
    para(title, size=24, color=NAVY, bold=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, after=14, line=1.05,
         style="Heading 1")
    p = para(blurb, size=11.5, color=SLATE, italic=True,
             align=WD_ALIGN_PARAGRAPH.CENTER, after=0, line=1.3, indent=0.9)
    p.paragraph_format.right_indent = Inches(0.9)


FIRST_OF_PART = [False]


def chap(num, title, lede_text, brk=None):
    """A chapter head.

    brk defaults to AUTOMATIC: only the first chapter after a part opener starts
    a fresh page, because the part opener already holds a page of its own.
    Forcing every chapter onto a new sheet is what produced the two-line widow
    pages, and in a compressed book it also costs real density. Pass brk=True to
    override where a chapter deserves its own opening.
    """
    if brk is None:
        brk = FIRST_OF_PART[0]
    FIRST_OF_PART[0] = False
    para(num, size=10.5, color=GOLD, bold=True, caps=True,
         align=WD_ALIGN_PARAGRAPH.LEFT, before=20, after=2, keep=True, brk=brk)
    q = para(title, size=15, color=NAVY, bold=True,
             align=WD_ALIGN_PARAGRAPH.LEFT, after=8, keep=True,
             style="Heading 2")
    _border(q)
    if lede_text:
        para(lede_text, size=10, color=SLATE, italic=True, before=4, after=10,
             align=WD_ALIGN_PARAGRAPH.LEFT, keep=True)


def sub(title):
    para(title, size=11, color=NAVY, bold=True, before=12, after=5,
         align=WD_ALIGN_PARAGRAPH.LEFT, keep=True, style="Heading 3")


def eq(t, size=12):
    para(t, size=size, color=NAVY, bold=True, font=MATH,
         align=WD_ALIGN_PARAGRAPH.CENTER, before=6, after=7, line=1.0)


def _shade(tc, fill):
    pr = tc.get_or_add_tcPr(); e = OxmlElement("w:shd")
    e.set(qn("w:val"), "clear"); e.set(qn("w:fill"), fill); pr.append(e)


def _mar(c, top=60, side=90, bot=60):
    pr = c._tc.get_or_add_tcPr(); m = OxmlElement("w:tcMar")
    for tag, w in (("top", top), ("start", side), ("bottom", bot), ("end", side)):
        e = OxmlElement("w:" + tag); e.set(qn("w:w"), str(w))
        e.set(qn("w:type"), "dxa"); m.append(e)
    pr.append(m)


def _borders(t, color="C7D3DE", sz=4):
    pr = t._tbl.tblPr; old = pr.find(qn("w:tblBorders"))
    if old is not None: pr.remove(old)
    b = OxmlElement("w:tblBorders")
    for tag in ("top", "start", "bottom", "end", "insideH", "insideV"):
        e = OxmlElement("w:" + tag)
        e.set(qn("w:val"), "single"); e.set(qn("w:sz"), str(sz))
        e.set(qn("w:color"), color); e.set(qn("w:space"), "0"); b.append(e)
    pr.append(b)


def _grid(t, widths):
    pr = t._tbl.tblPr
    old = pr.find(qn("w:tblLayout"))
    if old is not None: pr.remove(old)
    lay = OxmlElement("w:tblLayout"); lay.set(qn("w:type"), "fixed"); pr.append(lay)
    w = pr.find(qn("w:tblW"))
    if w is None: w = OxmlElement("w:tblW"); pr.append(w)
    w.set(qn("w:type"), "dxa"); w.set(qn("w:w"), str(int(sum(widths)*1440)))
    g = t._tbl.find(qn("w:tblGrid"))
    for gc in list(g): g.remove(gc)
    for wd in widths:
        gc = OxmlElement("w:gridCol"); gc.set(qn("w:w"), str(int(wd*1440)))
        g.append(gc)
    for row in t.rows:
        for j, c in enumerate(row.cells):
            if j < len(widths): c.width = Inches(widths[j])


def _cell(c, text, size=9, bold=False, color=INK, font=SERIF,
          align=WD_ALIGN_PARAGRAPH.LEFT, italic=False, raw=False):
    c.text = ""; p = c.paragraphs[0]
    p.paragraph_format.alignment = align; p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.12
    _runs(p, str(text), size, color, font, bold, italic, raw=raw)


def table(headers, rows, widths, size=9, mono=(), centre=(), colours=None,
          raw=False, hsize=8.5):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER; t.autofit = False
    _borders(t)
    for j, hh in enumerate(headers):
        c = t.rows[0].cells[j]; _shade(c._tc, HEADFILL); _mar(c, 70, bot=70)
        _cell(c, hh, size=hsize, bold=True, color=NAVY,
              align=WD_ALIGN_PARAGRAPH.CENTER)
    for i, row in enumerate(rows):
        cells = t.add_row().cells
        for j, v in enumerate(row):
            _mar(cells[j])
            col = INK
            if colours and (i, j) in colours: col = colours[(i, j)]
            _cell(cells[j], v, size=size, color=col,
                  font=MATH if j in mono else SERIF, raw=raw,
                  align=WD_ALIGN_PARAGRAPH.CENTER if j in centre
                  else WD_ALIGN_PARAGRAPH.LEFT)
    _grid(t, widths)
    tr = t.rows[0]._tr.get_or_add_trPr(); tr.append(OxmlElement("w:tblHeader"))
    for i, row in enumerate(t.rows):
        row._tr.get_or_add_trPr().append(OxmlElement("w:cantSplit"))
        if i < len(t.rows) - 1:
            for c in row.cells:
                for pp in c.paragraphs: pp.paragraph_format.keep_with_next = True
    para("", size=1, after=4)


def grade(kind, text):
    """The Ledger's status box, as a shaded two-cell strip."""
    fill, col = GRADEFILL.get(kind, (BOXFILL, SLATE))
    t = doc.add_table(rows=1, cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER; t.autofit = False
    _borders(t, color=fill, sz=2)
    lab, body = t.rows[0].cells
    _shade(lab._tc, fill); _shade(body._tc, fill)
    _mar(lab, 70, 95, 70); _mar(body, 70, 95, 70)
    _cell(lab, kind, size=8, bold=True, color=col,
          align=WD_ALIGN_PARAGRAPH.CENTER)
    _cell(body, text, size=9, color=INK)
    _grid(t, [0.98, TEXT_W - 0.98])
    t.rows[0]._tr.get_or_add_trPr().append(OxmlElement("w:cantSplit"))
    para("", size=1, after=5)


def callout(title, blocks, fill=BOXFILL, accent=RULECOL):
    para("", size=1, after=2)
    t = doc.add_table(rows=1, cols=1); t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False; _borders(t, color=fill, sz=2)
    c = t.rows[0].cells[0]; _shade(c._tc, fill); _mar(c, 150, 190, 150)
    _grid(t, [TEXT_W])
    t.rows[0]._tr.get_or_add_trPr().append(OxmlElement("w:cantSplit"))
    pr = c._tc.get_or_add_tcPr(); b = OxmlElement("w:tcBorders")
    e = OxmlElement("w:start"); e.set(qn("w:val"), "single")
    e.set(qn("w:sz"), "18"); e.set(qn("w:color"), accent)
    e.set(qn("w:space"), "0"); b.append(e); pr.append(b)
    c.text = ""; p = c.paragraphs[0]; p.paragraph_format.space_after = Pt(4)
    r = p.add_run(title.upper()); r.font.name = SERIF; r.font.size = Pt(9)
    r.font.color.rgb = GOLD; r.bold = True
    for i, blk in enumerate(blocks):
        q = c.add_paragraph()
        q.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        q.paragraph_format.space_after = Pt(0 if i == len(blocks)-1 else 6)
        q.paragraph_format.line_spacing = 1.18
        _runs(q, blk, 9.5, INK, SERIF)
    para("", size=1, after=4)


FIGN = [0]


def figure(fname, caption, width=6.5, brk=False):
    FIGN[0] += 1
    p = doc.add_paragraph(); pf = p.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf.space_before = Pt(7); pf.space_after = Pt(3); pf.keep_with_next = True
    if brk: pf.page_break_before = True
    p.add_run().add_picture(os.path.join(IMG, fname), width=Inches(width))
    q = doc.add_paragraph(); qf = q.paragraph_format
    qf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    qf.space_before = Pt(0); qf.space_after = Pt(9); qf.line_spacing = 1.10
    _runs(q, "Figure %d." % FIGN[0], 8.5, GOLD, SERIF, bold=True)
    _runs(q, "  " + caption, 8.5, SLATE, SERIF, italic=True)
    return FIGN[0]


# header / footer
hp = s.header.paragraphs[0]; hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
hr = hp.add_run("UNIFIED MECHANICS  ·  THE FIELD LEDGER  ·  SECOND EDITION")
hr.font.name = SERIF; hr.font.size = Pt(8); hr.font.color.rgb = GOLD
_border(hp, color="D8D8D8", sz=4)
fp = s.footer.paragraphs[0]; fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
fld = OxmlElement("w:fldSimple"); fld.set(qn("w:instr"), " PAGE ")
_r = OxmlElement("w:r"); _rp = OxmlElement("w:rPr")
_f = OxmlElement("w:rFonts"); _f.set(qn("w:ascii"), SERIF)
_f.set(qn("w:hAnsi"), SERIF); _rp.append(_f)
_sz = OxmlElement("w:sz"); _sz.set(qn("w:val"), "17"); _rp.append(_sz)
_c = OxmlElement("w:color"); _c.set(qn("w:val"), "5B6B7A"); _rp.append(_c)
_r.append(_rp); _t = OxmlElement("w:t"); _t.text = "1"; _r.append(_t)
fld.append(_r); fp._p.append(fld)
s.different_first_page_header_footer = True

# ── the certificates ──────────────────────────────────────────────────────
AD = r"C:\Users\joesh\Desktop\SCI\ADDENDA"
A09 = json.load(open(os.path.join(AD, "A09 The Master Equation",
                                  "master_equation_results.json")))
A03 = json.load(open(os.path.join(AD, "A03 The Forward Solver",
                                  "um_forward_solver_results.json")))
BY = {d["key"]: d for d in A09["rows"]}
ORDER = ["Omega_b", "Omega_c", "Omega_DE", "Y_He", "n_s", "tau"]
NAME = {"Omega_b": "Ω_b", "Omega_c": "Ω_c", "Omega_DE": "Ω_{DE}",
        "Y_He": "Y_{He}", "n_s": "n_s", "tau": "τ"}
CAND = A09["candidate_half_weight_typing"]
LAD = A09["typing_ladder"]
RUNG = {tuple(L["typed"]): L for L in LAD["rungs"]}
phi = (1 + math.sqrt(5)) / 2
r = 1 / (2 * phi); u = 1 - r; r3 = r ** 3
W_L, W_B, W_M = u * u, 2 * u * r, r * r
